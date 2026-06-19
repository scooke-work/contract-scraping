"""Offline unit tests for the core crawler helpers (no network)."""
import io

import pytest

import scrape_contracts as sc


def test_domain_of_strips_www_prefix_not_chars():
    assert sc.domain_of("https://www.example.com/x") == "example.com"
    # regression: lstrip("www.") wrongly turned "wto.org" into "to.org"
    assert sc.domain_of("https://wto.org/a") == "wto.org"
    assert sc.domain_of("https://www.law.cornell.edu/ucc/2") == "law.cornell.edu"
    assert sc.domain_of("https://EN.Wikipedia.ORG/wiki/X") == "en.wikipedia.org"
    assert sc.domain_of("https://example.com:8443/x") == "example.com"


def _rec(text, rid="x"):
    return {"id": rid, "source": "S", "category": "c", "url": "u", "title": "t", "text": text}


def test_chunks_respect_size_and_carry_overlap(tmp_path):
    cfg = sc.GlobalConfig(chunk_words=50, chunk_overlap=10, min_words=1)
    scraper = sc.ContractScraper(cfg, tmp_path)
    text = "\n".join(" ".join(f"w{i}_{j}" for j in range(20)) for i in range(12))
    chunks = list(scraper._chunks_for(_rec(text)))
    assert len(chunks) > 1
    for c in chunks:
        assert len(c["text"].split()) <= cfg.chunk_words
    assert [c["chunk_index"] for c in chunks] == list(range(len(chunks)))
    # overlap: the start of chunk 2 should repeat the tail of chunk 1
    assert chunks[1]["text"].split()[0] in chunks[0]["text"].split()


def test_chunks_hard_split_oversized_paragraph(tmp_path):
    cfg = sc.GlobalConfig(chunk_words=50, chunk_overlap=10, min_words=1)
    scraper = sc.ContractScraper(cfg, tmp_path)
    text = " ".join(f"w{i}" for i in range(140))  # one 140-word paragraph
    chunks = list(scraper._chunks_for(_rec(text)))
    assert len(chunks) >= 3
    for c in chunks:
        assert len(c["text"].split()) <= cfg.chunk_words


def test_edgar_doc_url_construction():
    hit = {"_id": "0001193125-21-000123:dex101.htm",
           "_source": {"ciks": ["0001232524"], "display_names": ["ACME (CIK 1)"]}}
    built = sc.ContractScraper._edgar_doc_url(hit)
    assert built is not None
    url, _src = built
    assert url == ("https://www.sec.gov/Archives/edgar/data/"
                   "1232524/000119312521000123/dex101.htm")


def test_edgar_doc_url_missing_fields_returns_none():
    assert sc.ContractScraper._edgar_doc_url({"_id": "acc", "_source": {}}) is None


def test_append_loads_existing_and_dedups(tmp_path):
    cfg = sc.GlobalConfig()
    first = sc.ContractScraper(cfg, tmp_path)
    rec = {"id": sc.url_id("http://a/1"), "source": "S", "category": "c",
           "url": "http://a/1", "title": "t", "text": "hello world", "word_count": 2}
    first._write_raw(rec)

    resumed = sc.ContractScraper(cfg, tmp_path, append=True)
    assert "http://a/1" in resumed.known_urls
    assert len(resumed.records) == 1


def test_docx_to_text_roundtrip():
    pytest.importorskip("docx")
    from docx import Document
    doc = Document()
    doc.add_paragraph("Master Services Agreement")
    doc.add_paragraph("This Agreement is governed by the laws of Delaware.")
    buf = io.BytesIO()
    doc.save(buf)
    text = sc.ContractScraper._docx_to_text(buf.getvalue())
    assert "Master Services Agreement" in text
    assert "Delaware" in text


def test_binary_skipped_when_disabled(tmp_path, monkeypatch):
    cfg = sc.GlobalConfig(parse_binaries=False)
    scraper = sc.ContractScraper(cfg, tmp_path)
    # stub the network: pretend the URL returned a PDF
    monkeypatch.setattr(scraper, "_fetch_bytes",
                        lambda url, headers=None: (b"%PDF-1.4 ...", "application/pdf"))
    assert scraper._fetch_and_extract_document("http://x/ex.pdf", {}) is None
