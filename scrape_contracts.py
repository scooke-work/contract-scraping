#!/usr/bin/env python3
"""
scrape_contracts.py
===================

A reusable, polite web scraper that builds an LLM-ingestible knowledge base
about Business-to-Business (B2B) contracts and contracting.

It reads a curated list of authoritative sources from ``sources.yaml``,
crawls them (respecting robots.txt and rate limits), extracts the clean
main text of each page, and writes the result in three LLM-friendly forms:

  1. output/raw/<source>/<id>.json   - one clean record per page (full text + metadata)
  2. output/chunks.jsonl             - one overlapping text chunk per line (ready for embedding / RAG)
  3. output/knowledge_base.md        - a single human-readable markdown digest
  plus output/manifest.json          - a summary of the run.

Usage
-----
    python scrape_contracts.py                      # crawl everything in sources.yaml
    python scrape_contracts.py --source "Cornell LII"   # only one source
    python scrape_contracts.py --max-pages 5        # cap pages per source (override)
    python scrape_contracts.py --list               # list configured sources and exit
    python scrape_contracts.py --output ./kb        # change output directory

See README.md for details.
"""

from __future__ import annotations

import argparse
import dataclasses
import hashlib
import io
import json
import re
import sys
import time
import urllib.robotparser
from collections import deque
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urljoin, urldefrag, urlparse

# ---- third-party (see requirements.txt) ------------------------------------
try:
    import requests
except ImportError:  # pragma: no cover
    sys.exit("Missing dependency 'requests'. Run: pip install -r requirements.txt")

try:
    import yaml
except ImportError:  # pragma: no cover
    sys.exit("Missing dependency 'PyYAML'. Run: pip install -r requirements.txt")

try:
    import trafilatura
except ImportError:  # pragma: no cover
    sys.exit("Missing dependency 'trafilatura'. Run: pip install -r requirements.txt")

from bs4 import BeautifulSoup  # bundled with trafilatura's deps; used for link discovery


# ---------------------------------------------------------------------------
# Configuration data classes
# ---------------------------------------------------------------------------
@dataclasses.dataclass
class GlobalConfig:
    user_agent: str = (
        "B2BContractKnowledgeBot/1.0 (+research; respects robots.txt)"
    )
    # SEC requires a descriptive UA with contact info for automated access.
    # Put a real email here before crawling EDGAR.
    edgar_user_agent: str = "B2BContractKnowledgeBot research ai@agreementsdemo.com"
    request_delay: float = 1.5          # seconds between requests to the same host
    timeout: int = 30                   # per-request timeout (s)
    max_pages_default: int = 25         # default cap per source if unset
    chunk_words: int = 800              # ~ target words per chunk
    chunk_overlap: int = 100            # words of overlap between chunks
    min_words: int = 60                 # skip pages with less clean text than this
    parse_binaries: bool = True         # EDGAR: parse PDF/Word exhibits to TEXT in-memory
    max_binary_mb: float = 25.0         # skip a binary exhibit larger than this (MB)


@dataclasses.dataclass
class Source:
    name: str
    category: str
    seeds: list[str] = dataclasses.field(default_factory=list)
    allowed_domains: list[str] = dataclasses.field(default_factory=list)
    follow_links: bool = False
    max_pages: int | None = None
    include_patterns: list[str] = dataclasses.field(default_factory=list)
    exclude_patterns: list[str] = dataclasses.field(default_factory=list)
    # --- EDGAR-only fields (used when kind == "edgar") -------------------
    kind: str = "web"                   # "web" (default) or "edgar"
    queries: list[str] = dataclasses.field(default_factory=list)
    forms: list[str] = dataclasses.field(default_factory=list)
    per_query: int | None = None        # max agreements to keep per query
    start_date: str | None = None       # EDGAR: only filings on/after YYYY-MM-DD
    end_date: str | None = None         # EDGAR: only filings on/before YYYY-MM-DD
    exhibit_types: list[str] = dataclasses.field(default_factory=list)
    # EDGAR: keep only hits whose exhibit type starts with one of these (e.g.
    # "EX-10" = material contracts). Empty = accept any document the hit points to.


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def slugify(text: str, maxlen: int = 60) -> str:
    text = re.sub(r"[^a-zA-Z0-9]+", "-", text.strip().lower()).strip("-")
    return text[:maxlen] or "untitled"


def url_id(url: str) -> str:
    return hashlib.sha1(url.encode("utf-8")).hexdigest()[:12]


def normalize_url(url: str) -> str:
    """Drop fragments and trailing whitespace; keep query strings."""
    url, _frag = urldefrag(url.strip())
    return url


def domain_of(url: str) -> str:
    # Strip the "www." prefix (and any :port). NB: do NOT use lstrip("www."),
    # which strips *characters* w/./ — e.g. "wto.org" -> "to.org".
    host = urlparse(url).netloc.lower().split(":", 1)[0]
    return host[4:] if host.startswith("www.") else host


def matches_any(url: str, patterns: list[str]) -> bool:
    return any(re.search(p, url) for p in patterns)


class RobotsCache:
    """Caches one RobotFileParser per host and answers can_fetch()."""

    def __init__(self, user_agent: str, session: requests.Session):
        self.user_agent = user_agent
        self.session = session
        self._cache: dict[str, urllib.robotparser.RobotFileParser | None] = {}

    def can_fetch(self, url: str) -> bool:
        parsed = urlparse(url)
        host = f"{parsed.scheme}://{parsed.netloc}"
        if host not in self._cache:
            rp = urllib.robotparser.RobotFileParser()
            robots_url = urljoin(host, "/robots.txt")
            try:
                resp = self.session.get(robots_url, timeout=15)
                if resp.status_code >= 400:
                    rp = None  # no robots / inaccessible -> allow
                else:
                    rp.parse(resp.text.splitlines())
            except requests.RequestException:
                rp = None
            self._cache[host] = rp
        rp = self._cache[host]
        if rp is None:
            return True
        try:
            return rp.can_fetch(self.user_agent, url)
        except Exception:
            return True


# ---------------------------------------------------------------------------
# Core crawler
# ---------------------------------------------------------------------------
class ContractScraper:
    def __init__(self, cfg: GlobalConfig, output_dir: Path, append: bool = False):
        self.cfg = cfg
        self.output_dir = output_dir
        self.session = requests.Session()
        self.session.headers.update({"User-Agent": cfg.user_agent})
        self.robots = RobotsCache(cfg.user_agent, self.session)
        self._last_hit: dict[str, float] = {}     # host -> last request time
        self.records: list[dict] = []             # all clean page records
        self.known_urls: set[str] = set()         # URLs already in the corpus (dedup)
        if append:
            self._load_existing()

    def _load_existing(self) -> None:
        """Append mode: read the per-page records already on disk back into memory
        so the rebuilt digest/chunks/manifest are the union of old + new, and so we
        skip any URL we already have (dedup-by-URL across runs)."""
        raw_dir = self.output_dir / "raw"
        if not raw_dir.exists():
            return
        loaded = 0
        for jf in raw_dir.rglob("*.json"):
            try:
                rec = json.loads(jf.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                continue
            if rec.get("url"):
                self.records.append(rec)
                self.known_urls.add(rec["url"])
                loaded += 1
        if loaded:
            print(f"[append] loaded {loaded} existing record(s); will skip URLs already collected")

    # -- politeness ---------------------------------------------------------
    def _throttle(self, url: str) -> None:
        host = urlparse(url).netloc
        last = self._last_hit.get(host, 0.0)
        wait = self.cfg.request_delay - (time.time() - last)
        if wait > 0:
            time.sleep(wait)
        self._last_hit[host] = time.time()

    def _fetch(self, url: str, extra_headers: dict | None = None) -> str | None:
        if not self.robots.can_fetch(url):
            print(f"    [robots] disallowed: {url}")
            return None
        self._throttle(url)
        try:
            resp = self.session.get(url, timeout=self.cfg.timeout, headers=extra_headers)
        except requests.RequestException as exc:
            print(f"    [error] {exc} :: {url}")
            return None
        if resp.status_code != 200:
            print(f"    [http {resp.status_code}] {url}")
            return None
        ctype = resp.headers.get("Content-Type", "")
        if "html" not in ctype and "xml" not in ctype:
            return None
        return resp.text

    # -- extraction ---------------------------------------------------------
    def _extract(self, html: str, url: str) -> dict | None:
        extracted = trafilatura.extract(
            html,
            url=url,
            include_comments=False,
            include_tables=True,
            favor_precision=True,
            output_format="json",
            with_metadata=True,
        )
        if not extracted:
            return None
        try:
            data = json.loads(extracted)
        except json.JSONDecodeError:
            return None
        text = (data.get("text") or "").strip()
        if len(text.split()) < self.cfg.min_words:
            return None
        return {
            "title": data.get("title") or "",
            "author": data.get("author") or "",
            "date": data.get("date") or "",
            "description": data.get("description") or "",
            "text": text,
        }

    def _discover_links(self, html: str, base_url: str, source: Source) -> list[str]:
        soup = BeautifulSoup(html, "html.parser")
        out: list[str] = []
        for a in soup.find_all("a", href=True):
            link = normalize_url(urljoin(base_url, a["href"]))
            if not link.startswith(("http://", "https://")):
                continue
            if domain_of(link) not in {domain_of(d) for d in source.allowed_domains}:
                continue
            if source.exclude_patterns and matches_any(link, source.exclude_patterns):
                continue
            if source.include_patterns and not matches_any(link, source.include_patterns):
                continue
            out.append(link)
        return out

    # -- per-source crawl ---------------------------------------------------
    def crawl_source(self, source: Source) -> int:
        if source.kind == "edgar":
            return self.crawl_edgar(source)
        cap = source.max_pages or self.cfg.max_pages_default
        print(f"\n=== {source.name}  ({source.category}) — up to {cap} pages ===")
        seen: set[str] = set()
        queue: deque[str] = deque(normalize_url(s) for s in source.seeds)
        collected = 0

        while queue and collected < cap:
            url = queue.popleft()
            if url in seen:
                continue
            seen.add(url)
            if url in self.known_urls:        # already collected in a prior run (--append)
                continue

            html = self._fetch(url)
            if html is None:
                continue

            rec = self._extract(html, url)
            if rec:
                record = {
                    "id": url_id(url),
                    "source": source.name,
                    "category": source.category,
                    "url": url,
                    "fetched_at": datetime.now(timezone.utc).isoformat(),
                    "word_count": len(rec["text"].split()),
                    **rec,
                }
                self.records.append(record)
                self.known_urls.add(url)
                self._write_raw(record)
                collected += 1
                print(f"    [{collected}/{cap}] {record['title'][:70] or url}")

            if source.follow_links and (collected + len(queue)) < cap * 3:
                for link in self._discover_links(html, url, source):
                    if link not in seen:
                        queue.append(link)

        print(f"--- {source.name}: collected {collected} page(s) ---")
        return collected

    # -- SEC EDGAR: real executed B2B contracts -----------------------------
    EDGAR_FTS = "https://efts.sec.gov/LATEST/search-index"

    def _get_json(self, url: str, params: dict, headers: dict) -> dict | None:
        self._throttle(url)
        try:
            resp = self.session.get(url, params=params, headers=headers,
                                    timeout=self.cfg.timeout)
        except requests.RequestException as exc:
            print(f"    [error] {exc} :: {url}")
            return None
        if resp.status_code != 200:
            print(f"    [http {resp.status_code}] {url}  params={params}")
            return None
        try:
            return resp.json()
        except ValueError:
            return None

    @staticmethod
    def _edgar_doc_url(hit: dict) -> tuple[str, dict] | None:
        """Build the public Archives URL for an exhibit from a search hit."""
        _id = hit.get("_id", "")
        src = hit.get("_source", {})
        if ":" in _id:
            accession, fname = _id.split(":", 1)
        else:
            accession, fname = _id, ""
        ciks = src.get("ciks") or []
        cik = (ciks[0] if ciks else "").lstrip("0")
        if not (cik and fname):
            return None
        acc = accession.replace("-", "")
        url = f"https://www.sec.gov/Archives/edgar/data/{cik}/{acc}/{fname}"
        return url, src

    # -- document fetch + extraction (HTML, or PDF/Word parsed to text) ------
    def _fetch_bytes(self, url: str, extra_headers: dict | None = None
                     ) -> tuple[bytes, str] | None:
        """Like _fetch, but returns raw bytes + content-type (no HTML-only filter)
        so the caller can decide how to handle PDFs / Word docs."""
        if not self.robots.can_fetch(url):
            print(f"    [robots] disallowed: {url}")
            return None
        self._throttle(url)
        try:
            resp = self.session.get(url, timeout=self.cfg.timeout, headers=extra_headers)
        except requests.RequestException as exc:
            print(f"    [error] {exc} :: {url}")
            return None
        if resp.status_code != 200:
            print(f"    [http {resp.status_code}] {url}")
            return None
        return resp.content, resp.headers.get("Content-Type", "").lower()

    def _fetch_and_extract_document(self, url: str, headers: dict) -> dict | None:
        """Fetch an EDGAR exhibit and return a clean record dict. HTML exhibits go
        through trafilatura; PDF/Word exhibits are parsed to TEXT in memory (only
        when cfg.parse_binaries is on). The source file is NEVER written to disk —
        we keep the extracted text + metadata only."""
        got = self._fetch_bytes(url, headers)
        if got is None:
            return None
        content, ctype = got
        low = url.lower()
        is_pdf = ("application/pdf" in ctype) or low.endswith(".pdf")
        is_docx = ("wordprocessingml" in ctype) or low.endswith(".docx")
        if is_pdf or is_docx:
            kind = "PDF" if is_pdf else "DOCX"
            if not self.cfg.parse_binaries:
                print(f"    [skip {kind}] (parse_binaries off): {url}")
                return None
            if len(content) > self.cfg.max_binary_mb * 1024 * 1024:
                print(f"    [skip {kind} >{self.cfg.max_binary_mb:g}MB] {url}")
                return None
            text = self._pdf_to_text(content) if is_pdf else self._docx_to_text(content)
            if len(text.split()) < max(self.cfg.min_words, 100):
                return None
            print(f"    [parsed {kind} in-memory] {len(text.split())} words")
            return {"title": "", "author": "", "date": "", "description": "",
                    "text": text.strip()}
        # HTML / XML / plain text
        html = content.decode("utf-8", errors="replace")
        return self._extract_edgar(html, url)

    @staticmethod
    def _pdf_to_text(data: bytes) -> str:
        try:
            from pdfminer.high_level import extract_text   # lazy: optional dep
        except ImportError:
            print("    [pdf] pdfminer.six not installed — run: pip install pdfminer.six")
            return ""
        try:
            return (extract_text(io.BytesIO(data)) or "").strip()
        except Exception as exc:                           # malformed PDF, etc.
            print(f"    [pdf error] {exc}")
            return ""

    @staticmethod
    def _docx_to_text(data: bytes) -> str:
        try:
            import docx                                     # lazy: optional dep (python-docx)
        except ImportError:
            print("    [docx] python-docx not installed — run: pip install python-docx")
            return ""
        try:
            doc = docx.Document(io.BytesIO(data))
        except Exception as exc:
            print(f"    [docx error] {exc}")
            return ""
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    def _extract_edgar(self, html: str, url: str) -> dict | None:
        """Contracts are long and irregular; favour recall, then fall back."""
        text = trafilatura.extract(
            html, url=url, include_comments=False, include_tables=True,
            favor_precision=False, no_fallback=False,
        ) or ""
        title = ""
        if len(text.split()) < 150:
            soup = BeautifulSoup(html, "html.parser")
            if soup.title and soup.title.string:
                title = soup.title.string.strip()
            bs_text = soup.get_text(" ", strip=True)
            if len(bs_text.split()) > len(text.split()):
                text = bs_text
        if len(text.split()) < max(self.cfg.min_words, 100):
            return None
        return {"title": title, "author": "", "date": "", "description": "",
                "text": text.strip()}

    def crawl_edgar(self, source: Source) -> int:
        cap = source.max_pages or self.cfg.max_pages_default
        queries = source.queries or [source.name]
        per_query = source.per_query or max(1, cap // len(queries))
        headers = {"User-Agent": self.cfg.edgar_user_agent,
                   "Accept": "application/json"}
        doc_headers = {"User-Agent": self.cfg.edgar_user_agent}
        print(f"\n=== {source.name}  ({source.category}) — EDGAR full-text search, "
              f"up to {cap} agreements ===")
        seen: set[str] = set()
        collected = 0

        for q in queries:
            if collected >= cap:
                break
            print(f"  query: \"{q}\"")
            got, frm = 0, 0
            while got < per_query and collected < cap:
                params = {"q": f'"{q}"', "from": frm}
                if source.forms:
                    params["forms"] = ",".join(source.forms)
                if source.start_date or source.end_date:
                    params["dateRange"] = "custom"
                    if source.start_date:
                        params["startdt"] = source.start_date
                    if source.end_date:
                        params["enddt"] = source.end_date
                data = self._get_json(self.EDGAR_FTS, params, headers)
                hits = (data or {}).get("hits", {}).get("hits", [])
                if not hits:
                    break
                for hit in hits:
                    if collected >= cap or got >= per_query:
                        break
                    if hit.get("_id") in seen:
                        continue
                    seen.add(hit.get("_id"))
                    built = self._edgar_doc_url(hit)
                    if not built:
                        continue
                    url, meta = built
                    if source.exhibit_types:             # keep only real contract exhibits
                        ftype = (meta.get("file_type") or "").upper()
                        if not any(ftype.startswith(t.upper()) for t in source.exhibit_types):
                            continue
                    if url in self.known_urls:           # already collected (--append)
                        continue
                    rec = self._fetch_and_extract_document(url, doc_headers)
                    if not rec:
                        continue
                    names = meta.get("display_names") or []
                    record = {
                        "id": url_id(url),
                        "source": source.name,
                        "category": source.category,
                        "url": url,
                        "company": names[0] if names else "",
                        "form": meta.get("file_type") or meta.get("root_form") or "",
                        "file_date": meta.get("file_date") or "",
                        "matched_query": q,
                        "fetched_at": datetime.now(timezone.utc).isoformat(),
                        "word_count": len(rec["text"].split()),
                        **rec,
                    }
                    if not record["title"]:
                        record["title"] = f"{q.title()} — {record['company']}".strip(" —")
                    self.records.append(record)
                    self.known_urls.add(url)
                    self._write_raw(record)
                    collected += 1
                    got += 1
                    print(f"    [{collected}/{cap}] {record['title'][:70]}")
                frm += len(hits)
                if len(hits) < 10:
                    break

        print(f"--- {source.name}: collected {collected} agreement(s) ---")
        return collected

    # -- output writers -----------------------------------------------------
    def _write_raw(self, record: dict) -> None:
        folder = self.output_dir / "raw" / slugify(record["source"])
        folder.mkdir(parents=True, exist_ok=True)
        fname = f"{record['id']}-{slugify(record['title'] or 'page', 40)}.json"
        (folder / fname).write_text(
            json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8"
        )

    def _chunks_for(self, record: dict):
        """Pack paragraphs (split on blank lines) up to ~chunk_words, carrying a
        small word overlap between chunks. Avoids cutting mid-paragraph; falls
        back to fixed word windows for any single oversized paragraph."""
        size, overlap = self.cfg.chunk_words, self.cfg.chunk_overlap
        # trafilatura separates paragraphs with single newlines, so split on any
        # run of newlines (not just blank lines) to chunk on real text boundaries.
        paragraphs = [p.strip() for p in re.split(r"\n+", record["text"]) if p.strip()]
        if not paragraphs:
            paragraphs = [record["text"].strip()]

        chunks: list[list[str]] = []
        cur: list[str] = []
        for para in paragraphs:
            pw = para.split()
            if not pw:
                continue
            if len(pw) > size:                       # oversized paragraph: hard-split it
                if cur:
                    chunks.append(cur)
                    cur = []
                step = max(size - overlap, 1)
                for s in range(0, len(pw), step):
                    chunks.append(pw[s:s + size])
                    if s + size >= len(pw):
                        break
                continue
            if cur and len(cur) + len(pw) > size:    # would overflow: close current chunk
                chunks.append(cur)
                cur = cur[-overlap:] if overlap else []   # carry overlap
            cur.extend(pw)
        if cur:
            chunks.append(cur)

        for idx, words in enumerate(chunks):
            yield {
                "chunk_id": f"{record['id']}-{idx}",
                "source": record["source"],
                "category": record["category"],
                "url": record["url"],
                "title": record["title"],
                "chunk_index": idx,
                "text": " ".join(words),
            }

    def write_outputs(self) -> dict:
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # chunks.jsonl
        chunk_count = 0
        with (self.output_dir / "chunks.jsonl").open("w", encoding="utf-8") as fh:
            for rec in self.records:
                for chunk in self._chunks_for(rec):
                    fh.write(json.dumps(chunk, ensure_ascii=False) + "\n")
                    chunk_count += 1

        # knowledge_base.md
        by_source: dict[str, list[dict]] = {}
        for rec in self.records:
            by_source.setdefault(rec["source"], []).append(rec)

        md = ["# B2B Contracts Knowledge Base",
              f"\n_Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} "
              f"from {len(self.records)} pages across {len(by_source)} sources._\n"]
        for src, recs in sorted(by_source.items()):
            md.append(f"\n---\n\n## {src}\n")
            for rec in recs:
                md.append(f"\n### {rec['title'] or rec['url']}\n")
                md.append(f"_Source: [{rec['url']}]({rec['url']}) — "
                          f"category: {rec['category']} — {rec['word_count']} words_\n")
                md.append(rec["text"] + "\n")
        (self.output_dir / "knowledge_base.md").write_text("\n".join(md), encoding="utf-8")

        # manifest.json
        manifest = {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "pages": len(self.records),
            "chunks": chunk_count,
            "sources": {
                src: {
                    "pages": len(recs),
                    "words": sum(r["word_count"] for r in recs),
                    "category": recs[0]["category"],
                }
                for src, recs in by_source.items()
            },
        }
        (self.output_dir / "manifest.json").write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        return manifest


# ---------------------------------------------------------------------------
# Config loading
# ---------------------------------------------------------------------------
def load_config(path: Path) -> tuple[GlobalConfig, list[Source]]:
    data = yaml.safe_load(path.read_text(encoding="utf-8"))
    cfg = GlobalConfig(**(data.get("config") or {}))
    sources = []
    for s in data.get("sources", []):
        seeds = s.get("seeds", [])
        allowed = s.get("allowed_domains") or ([domain_of(seeds[0])] if seeds else [])
        sources.append(Source(
            name=s["name"],
            category=s.get("category", "general"),
            seeds=seeds,
            allowed_domains=allowed,
            follow_links=s.get("follow_links", False),
            max_pages=s.get("max_pages"),
            include_patterns=s.get("include_patterns", []),
            exclude_patterns=s.get("exclude_patterns", []),
            kind=s.get("type", "web"),
            queries=s.get("queries", []),
            forms=s.get("forms", []),
            per_query=s.get("per_query"),
            start_date=s.get("start_date"),
            end_date=s.get("end_date"),
            exhibit_types=s.get("exhibit_types", []),
        ))
    return cfg, sources


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Scrape authoritative sources into a B2B-contracts knowledge base.")
    ap.add_argument("--sources", default="sources.yaml", help="Path to sources.yaml")
    ap.add_argument("--output", default="output", help="Output directory")
    ap.add_argument("--source", action="append", help="Only crawl source(s) with this name (repeatable)")
    ap.add_argument("--max-pages", type=int, help="Override per-source page cap")
    ap.add_argument("--delay", type=float, help="Override per-host request delay (seconds)")
    ap.add_argument("--append", action="store_true",
                    help="Merge into an existing --output dir (dedup by URL) instead of rebuilding from scratch")
    ap.add_argument("--no-binaries", action="store_true",
                    help="Skip PDF/Word exhibits instead of parsing their text in-memory")
    ap.add_argument("--list", action="store_true", help="List configured sources and exit")
    args = ap.parse_args(argv)

    cfg_path = Path(args.sources)
    if not cfg_path.exists():
        sys.exit(f"Config not found: {cfg_path}")
    cfg, sources = load_config(cfg_path)
    if args.delay is not None:
        cfg.request_delay = args.delay

    if args.list:
        print("Configured sources:")
        for s in sources:
            cap = s.max_pages or cfg.max_pages_default
            if s.kind == "edgar":
                print(f"  - {s.name}  [{s.category}]  EDGAR  "
                      f"queries={len(s.queries)}  cap={cap}")
            else:
                print(f"  - {s.name}  [{s.category}]  seeds={len(s.seeds)}  "
                      f"follow={s.follow_links}  cap={cap}")
        return 0

    if args.source:
        wanted = {n.lower() for n in args.source}
        sources = [s for s in sources if s.name.lower() in wanted]
        if not sources:
            sys.exit("No matching sources. Use --list to see names.")
    if args.max_pages:
        for s in sources:
            s.max_pages = args.max_pages

    if args.no_binaries:
        cfg.parse_binaries = False

    scraper = ContractScraper(cfg, Path(args.output), append=args.append)
    total = 0
    for src in sources:
        total += scraper.crawl_source(src)

    if total == 0 and not scraper.records:
        print("\nNo pages collected. Check connectivity, robots rules, or your seeds.")
        return 1
    if total == 0:
        print("\nNo new pages collected; rewriting outputs from the existing corpus.")

    manifest = scraper.write_outputs()
    print("\n=========================================")
    print(f"Done. {manifest['pages']} pages -> {manifest['chunks']} chunks.")
    print(f"Output written to: {Path(args.output).resolve()}")
    print("  - chunks.jsonl        (LLM / RAG ingestion)")
    print("  - knowledge_base.md   (human-readable digest)")
    print("  - raw/<source>/*.json (per-page records)")
    print("  - manifest.json       (run summary)")
    print("=========================================")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
